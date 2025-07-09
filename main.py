import pandas as pd
from docx import Document
data = pd.read_csv("student_data.csv")
data_score = data["Score"]

number_of_student = len(data_score)
list_of_score = data_score.to_list()

mean_score = data_score.mean()
median_score = data_score.median()
mode_score = data_score.mode().tolist()
max_score = data_score.max()
min_score = data_score.min()
standard = data_score.std()

passed_student = (data_score >= 50).sum()
failed_student = (data_score < 50).sum()
report=f'''
Total Students: {number_of_student}
Median Score: {median_score}
Mean Score: {mean_score}
Mode Score(s):, {", ".join(map(str, mode_score))}
Max Score:, {max_score}
Min Score:, {min_score}
Standard Deviation:, {standard}
Passed Students:, {passed_student}
Failed Students:, {failed_student}
'''
doc = Document()
doc.add_heading("Student Score Report",level=1)
doc.add_paragraph(report)
doc.save('Student_score.docx')